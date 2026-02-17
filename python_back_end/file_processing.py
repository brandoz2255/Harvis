import base64
import io
import logging
import tempfile
import os
from typing import Optional, List, Tuple

# Configure logger
logger = logging.getLogger(__name__)

# Maximum pages to convert (to avoid memory issues with large documents)
MAX_PAGES_FOR_VISION = int(os.environ.get("MAX_PAGES_FOR_VISION", "10"))

def extract_text_from_file(file_data: str, file_type: str) -> Optional[str]:
    """
    Extract text from a file based on its type.
    
    Args:
        file_data: Base64 encoded file data
        file_type: MIME type or extension of the file
        
    Returns:
        Extracted text or None if extraction failed or type not supported
    """
    try:
        # Decode base64 data
        if ',' in file_data:
            _, encoded = file_data.split(',', 1)
        else:
            encoded = file_data
            
        decoded_data = base64.b64decode(encoded)
        file_bytes = io.BytesIO(decoded_data)
        
        # Determine extraction method based on type
        if 'pdf' in file_type.lower():
            return _extract_from_pdf(file_bytes)
        elif 'word' in file_type.lower() or 'docx' in file_type.lower():
            return _extract_from_docx(file_bytes)
        elif 'text' in file_type.lower() or 'txt' in file_type.lower() or 'md' in file_type.lower() or 'csv' in file_type.lower():
            return decoded_data.decode('utf-8', errors='replace')
        else:
            logger.warning(f"Unsupported file type for text extraction: {file_type}")
            return None
            
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return f"[Error extracting text from file: {str(e)}]"

def _extract_from_pdf(file_bytes: io.BytesIO) -> str:
    """Extract text from PDF BytesIO object."""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_bytes)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return "\n".join(text)
    except ImportError:
        logger.error("pypdf is not installed")
        return "[Error: pypdf library is missing]"
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise

def _extract_from_docx(file_bytes: io.BytesIO) -> str:
    """Extract text from DOCX BytesIO object."""
    try:
        import docx
        doc = docx.Document(file_bytes)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except ImportError:
        logger.error("python-docx is not installed")
        return "[Error: python-docx library is missing]"
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise


# ============================================================================
# Vision Model Support - Convert documents to images
# ============================================================================

def convert_file_to_images(file_data: str, file_type: str, max_pages: int = None) -> List[Tuple[str, str]]:
    """
    Convert a document file to images for vision model processing.

    Args:
        file_data: Base64 encoded file data
        file_type: MIME type or extension of the file
        max_pages: Maximum number of pages to convert (defaults to MAX_PAGES_FOR_VISION)

    Returns:
        List of tuples: [(base64_image, mime_type), ...]
        Returns empty list if conversion fails or type not supported
    """
    if max_pages is None:
        max_pages = MAX_PAGES_FOR_VISION

    try:
        # Decode base64 data
        if ',' in file_data:
            _, encoded = file_data.split(',', 1)
        else:
            encoded = file_data

        decoded_data = base64.b64decode(encoded)
        file_bytes = io.BytesIO(decoded_data)

        # Determine conversion method based on type
        file_type_lower = file_type.lower()

        if 'pdf' in file_type_lower:
            return _pdf_to_images(file_bytes, max_pages)
        elif 'word' in file_type_lower or 'docx' in file_type_lower or file_type_lower.endswith('.docx'):
            return _docx_to_images(file_bytes, max_pages)
        elif any(img_type in file_type_lower for img_type in ['image', 'png', 'jpg', 'jpeg', 'gif', 'webp']):
            # Already an image, just return it
            mime = 'image/png' if 'png' in file_type_lower else 'image/jpeg'
            return [(encoded, mime)]
        else:
            logger.warning(f"Unsupported file type for image conversion: {file_type}")
            return []

    except Exception as e:
        logger.error(f"Error converting file to images: {e}")
        return []


def _pdf_to_images(file_bytes: io.BytesIO, max_pages: int) -> List[Tuple[str, str]]:
    """
    Convert PDF pages to images using PyMuPDF (fitz).

    Returns list of (base64_image, mime_type) tuples.
    """
    images = []

    try:
        import fitz  # PyMuPDF

        # Open PDF from bytes
        pdf_document = fitz.open(stream=file_bytes.read(), filetype="pdf")
        total_pages = len(pdf_document)
        pages_to_convert = min(total_pages, max_pages)

        logger.info(f"Converting {pages_to_convert}/{total_pages} PDF pages to images")

        for page_num in range(pages_to_convert):
            page = pdf_document[page_num]

            # Render page to image (2x zoom for better quality)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PNG bytes
            img_bytes = pix.tobytes("png")
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')

            images.append((img_base64, "image/png"))
            logger.debug(f"Converted PDF page {page_num + 1}/{pages_to_convert}")

        pdf_document.close()
        logger.info(f"Successfully converted {len(images)} PDF pages to images")
        return images

    except ImportError:
        logger.error("PyMuPDF (fitz) is not installed. Install with: pip install PyMuPDF")
        # Fallback: try pdf2image
        return _pdf_to_images_fallback(file_bytes, max_pages)
    except Exception as e:
        logger.error(f"PDF to image conversion error: {e}")
        return []


def _pdf_to_images_fallback(file_bytes: io.BytesIO, max_pages: int) -> List[Tuple[str, str]]:
    """
    Fallback PDF to image conversion using pdf2image (requires poppler).
    """
    images = []

    try:
        from pdf2image import convert_from_bytes

        file_bytes.seek(0)
        pdf_bytes = file_bytes.read()

        # Convert PDF to images
        pil_images = convert_from_bytes(
            pdf_bytes,
            first_page=1,
            last_page=max_pages,
            dpi=150  # Good balance of quality vs size
        )

        logger.info(f"Converting {len(pil_images)} PDF pages using pdf2image")

        for i, pil_img in enumerate(pil_images):
            # Convert PIL image to base64 PNG
            img_buffer = io.BytesIO()
            pil_img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            img_base64 = base64.b64encode(img_buffer.read()).decode('utf-8')

            images.append((img_base64, "image/png"))
            logger.debug(f"Converted PDF page {i + 1}/{len(pil_images)}")

        return images

    except ImportError:
        logger.error("Neither PyMuPDF nor pdf2image is installed for PDF conversion")
        return []
    except Exception as e:
        logger.error(f"PDF to image fallback error: {e}")
        return []


def _docx_to_images(file_bytes: io.BytesIO, max_pages: int) -> List[Tuple[str, str]]:
    """
    Convert DOCX to images.

    Strategy:
    1. Try to convert DOCX to PDF using LibreOffice (if available)
    2. Then convert PDF pages to images
    3. Fallback: Extract text and create a simple image representation
    """
    images = []

    try:
        import subprocess
        import shutil

        # Check if LibreOffice is available
        libreoffice_path = shutil.which('libreoffice') or shutil.which('soffice')

        if libreoffice_path:
            return _docx_to_images_via_libreoffice(file_bytes, max_pages, libreoffice_path)
        else:
            logger.warning("LibreOffice not found, using text-based fallback for DOCX")
            return _docx_to_images_fallback(file_bytes, max_pages)

    except Exception as e:
        logger.error(f"DOCX to image conversion error: {e}")
        return _docx_to_images_fallback(file_bytes, max_pages)


def _docx_to_images_via_libreoffice(file_bytes: io.BytesIO, max_pages: int, libreoffice_path: str) -> List[Tuple[str, str]]:
    """
    Convert DOCX to images using LibreOffice.
    """
    import subprocess

    with tempfile.TemporaryDirectory() as temp_dir:
        # Save DOCX to temp file
        docx_path = os.path.join(temp_dir, "document.docx")
        with open(docx_path, 'wb') as f:
            file_bytes.seek(0)
            f.write(file_bytes.read())

        # Convert DOCX to PDF using LibreOffice
        try:
            subprocess.run(
                [libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, docx_path],
                capture_output=True,
                timeout=60,
                check=True
            )
        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out")
            return _docx_to_images_fallback(file_bytes, max_pages)
        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: {e}")
            return _docx_to_images_fallback(file_bytes, max_pages)

        # Find the generated PDF
        pdf_path = os.path.join(temp_dir, "document.pdf")
        if not os.path.exists(pdf_path):
            logger.error("LibreOffice did not generate PDF")
            return _docx_to_images_fallback(file_bytes, max_pages)

        # Convert PDF to images
        with open(pdf_path, 'rb') as f:
            pdf_bytes = io.BytesIO(f.read())

        return _pdf_to_images(pdf_bytes, max_pages)


def _docx_to_images_fallback(file_bytes: io.BytesIO, max_pages: int) -> List[Tuple[str, str]]:
    """
    Fallback: Create a simple text-based image from DOCX content.
    Uses PIL to render text onto an image.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        import docx

        # Extract text from DOCX
        file_bytes.seek(0)
        doc = docx.Document(file_bytes)

        # Collect paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        if not paragraphs:
            logger.warning("No text content found in DOCX")
            return []

        # Create images (one per "page" of text)
        images = []

        # Image dimensions (A4-ish aspect ratio)
        img_width = 800
        img_height = 1100
        margin = 50
        line_height = 20
        lines_per_page = (img_height - 2 * margin) // line_height

        # Try to get a font
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
        except:
            font = ImageFont.load_default()

        # Split text into lines
        all_lines = []
        for para in paragraphs:
            # Word wrap
            words = para.split()
            current_line = ""
            for word in words:
                test_line = f"{current_line} {word}".strip()
                if len(test_line) > 80:  # Approximate line width
                    if current_line:
                        all_lines.append(current_line)
                    current_line = word
                else:
                    current_line = test_line
            if current_line:
                all_lines.append(current_line)
            all_lines.append("")  # Paragraph break

        # Create pages
        page_count = 0
        for i in range(0, len(all_lines), lines_per_page):
            if page_count >= max_pages:
                break

            page_lines = all_lines[i:i + lines_per_page]

            # Create image
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)

            # Draw text
            y = margin
            for line in page_lines:
                draw.text((margin, y), line, fill='black', font=font)
                y += line_height

            # Add page indicator
            draw.text(
                (img_width - margin - 50, img_height - margin),
                f"Page {page_count + 1}",
                fill='gray',
                font=font
            )

            # Convert to base64
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            img_base64 = base64.b64encode(img_buffer.read()).decode('utf-8')

            images.append((img_base64, "image/png"))
            page_count += 1

        logger.info(f"Created {len(images)} images from DOCX text content")
        return images

    except ImportError as e:
        logger.error(f"Missing library for DOCX fallback: {e}")
        return []
    except Exception as e:
        logger.error(f"DOCX fallback conversion error: {e}")
        return []


def is_vision_compatible_file(file_type: str) -> bool:
    """
    Check if a file type can be converted to images for vision models.
    """
    file_type_lower = file_type.lower()
    compatible_types = [
        'pdf', 'docx', 'doc', 'word',
        'image', 'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'
    ]
    return any(t in file_type_lower for t in compatible_types)
