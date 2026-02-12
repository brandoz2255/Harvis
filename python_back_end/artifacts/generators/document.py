"""
Word document generator using python-docx
"""

import os
import uuid
import logging
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


def generate_document(content: Dict[str, Any], output_dir: str) -> str:
    """
    Generate DOCX file from manifest content.

    Args:
        content: Document content dict with 'sections' key
        output_dir: Directory to save the file

    Returns:
        Full path to generated file

    Expected content format:
    {
        "sections": [
            {"type": "heading", "level": 1, "content": "Title"},
            {"type": "paragraph", "content": "Text content", "bold": false, "italic": false},
            {"type": "table", "rows": [["Header1", "Header2"], ["Data1", "Data2"]]},
            {"type": "list", "items": ["Item 1", "Item 2"]},
            {"type": "quote", "content": "Quoted text"},
            {"type": "code", "content": "print('hello')"}
        ],
        "title": "Document Title",
        "author": "Harvis AI"
    }
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Set document properties
    core_properties = doc.core_properties
    core_properties.author = content.get("author", "Harvis AI")
    if content.get("title"):
        core_properties.title = content["title"]
    if content.get("subject"):
        core_properties.subject = content["subject"]

    sections = content.get("sections", [])

    for section in sections:
        section_type = section.get("type", "paragraph")

        if section_type == "heading":
            level = min(max(section.get("level", 1), 0), 9)
            heading = doc.add_heading(section.get("content", ""), level=level)

        elif section_type == "paragraph":
            para = doc.add_paragraph()
            run = para.add_run(section.get("content", ""))

            if section.get("bold"):
                run.bold = True
            if section.get("italic"):
                run.italic = True
            if section.get("underline"):
                run.underline = True

        elif section_type == "table":
            rows = section.get("rows", [])
            if rows:
                num_rows = len(rows)
                num_cols = max(len(row) for row in rows) if rows else 0

                if num_cols > 0:
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'

                    for i, row in enumerate(rows):
                        for j, cell_value in enumerate(row):
                            if j < num_cols:
                                table.cell(i, j).text = str(cell_value)

                                # Make first row bold (headers)
                                if i == 0:
                                    for paragraph in table.cell(i, j).paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                    doc.add_paragraph()  # Add spacing after table

        elif section_type == "list":
            items = section.get("items", [])
            list_style = section.get("style", "bullet")  # bullet or numbered

            for item in items:
                if list_style == "numbered":
                    doc.add_paragraph(str(item), style='List Number')
                else:
                    doc.add_paragraph(str(item), style='List Bullet')

        elif section_type == "quote":
            para = doc.add_paragraph()
            para.style = 'Quote'
            run = para.add_run(section.get("content", ""))
            run.italic = True

        elif section_type == "code":
            para = doc.add_paragraph()
            run = para.add_run(section.get("content", ""))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            # Add light gray background effect via formatting
            para.paragraph_format.left_indent = Inches(0.5)

        elif section_type == "image":
            # Image from URL or path - placeholder for now
            image_path = section.get("path") or section.get("url")
            if image_path and os.path.exists(image_path):
                width = section.get("width")
                if width:
                    doc.add_picture(image_path, width=Inches(width))
                else:
                    doc.add_picture(image_path, width=Inches(6))

        elif section_type == "page_break":
            doc.add_page_break()

    # Generate unique filename
    filename = f"artifact_{uuid.uuid4().hex[:16]}.docx"
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)
    logger.info(f"Generated document: {filepath}")

    return filepath
